#!/usr/bin/env python3
"""
Test file to generate images using Google's Nano Banana Pro API (gemini-3-pro-image-preview)
and save the generated images to a DOCX file.
Based on the official Google Gemini API documentation for image generation.
"""

import json
import tempfile
import os
from pathlib import Path

# Import the Google GenAI SDK
try:
    from google import genai
    from google.genai import types
    GENAI_AVAILABLE = True
except ImportError:
    print("google-genai not available. Please install it with: pip install google-genai")
    GENAI_AVAILABLE = False

# Attempt to import docx, with fallback if not available
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    print("python-docx not available. Please install it with: pip install python-docx")
    DOCX_AVAILABLE = False


def load_api_keys():
    """Load API keys from key.json file."""
    try:
        with open('key.json', 'r') as f:
            data = json.load(f)

        apis = {api['name']: api['keys'][0] for api in data['apis']}
        return apis
    except FileNotFoundError:
        print("key.json file not found!")
        return {}
    except Exception as e:
        print(f"Error loading API keys: {e}")
        return {}


def generate_image_with_nanobanana_pro(prompt, api_key, aspect_ratio="16:9", image_size="2K"):
    """
    Generate an image using the Gemini 3 Pro Image Preview model via the SDK.
    """
    if not GENAI_AVAILABLE:
        return False, "Google GenAI SDK not available. Please install it with: pip install google-genai"

    if not api_key or "YOUR_" in api_key:
        return False, "No valid API key provided."

    try:
        client = genai.Client(api_key=api_key)

        # Enhance prompt with aspect ratio since we removed the config
        # (Many preview models prioritize prompt instructions for dimensions)
        full_prompt = f"{prompt} --aspect_ratio {aspect_ratio}"

        # Call the API without the problematic response_mime_type config
        response = client.models.generate_content(
            model="gemini-3-pro-image-preview",
            contents=[full_prompt],
            # If you need safety settings, add them here, but do NOT add response_mime_type="image/jpeg"
        )

        # Check for image data in the response
        if response.candidates:
            for candidate in response.candidates:
                if hasattr(candidate, 'content') and candidate.content.parts:
                    for part in candidate.content.parts:
                        # Check for inline_data (image)
                        if hasattr(part, 'inline_data') and part.inline_data:
                            # Create a temporary file
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as temp_file:
                                temp_file.write(part.inline_data.data)
                                return True, temp_file.name

        # If no image found, check for text feedback (e.g., safety refusal)
        if response.candidates and response.candidates[0].content.parts:
            text_part = response.candidates[0].content.parts[0].text
            if text_part:
                return False, f"Model returned text instead of image: {text_part}"

        return False, "No image found in response. (Check safety ratings or prompt)"

    except Exception as e:
        return False, f"SDK Error: {str(e)}"


def save_images_to_docx(images_data, output_filename="generated_images.docx"):
    """
    Save generated images to a DOCX file with proper formatting.

    Args:
        images_data (list): List of tuples containing (prompt, image_path)
        output_filename (str): Name of the output DOCX file

    Returns:
        str: Success or error message
    """
    if not DOCX_AVAILABLE:
        return "DOCX functionality not available. Please install python-docx."

    try:
        # Create a new document
        doc = Document()
        
        # Add title
        doc.add_heading('Generated Images Report', 0)
        
        # Add a paragraph with description
        doc.add_paragraph(
            'This document contains images generated using Google\'s Nano Banana Pro API '
            '(gemini-3-pro-image-preview) based on the provided prompts.'
        )
        
        # Add each image with its prompt
        for i, (prompt, image_path) in enumerate(images_data):
            doc.add_heading(f'Image {i+1}: "{prompt}"', level=1)
            
            # Add the image to the document
            if os.path.exists(image_path):
                # Adjust width as needed - using 6 inches as default
                doc.add_picture(image_path, width=Inches(6))
                
                # Add a blank paragraph for spacing
                doc.add_paragraph()
            else:
                doc.add_paragraph(f"Image file not found: {image_path}")
        
        # Save the document
        doc.save(output_filename)
        return f"Successfully saved images to {output_filename}"
        
    except Exception as e:
        return f"Error saving DOCX: {e}"


def main():
    """Main function to test image generation and save to DOCX."""
    print("Testing Image Generation with Nano Banana Pro API and DOCX Output")
    print("=" * 70)

    # Load API keys
    api_keys = load_api_keys()

    if not api_keys:
        print("No API keys found. Please ensure key.json exists with valid API keys.")
        return

    # Get the Nano Banana Pro API key
    nano_banana_key = api_keys.get('nano_banana_gemini')
    if not nano_banana_key:
        print("Nano Banana Pro API key not found in key.json")
        return

    print(f"Found Nano Banana Pro API key: {nano_banana_key[:10]}...")
    print("")

    # Define test prompts - simpler prompts for better compatibility
    prompts = [
        "A beautiful sunset over mountains",
        "A futuristic city skyline at night",
        "A red rose flower with dewdrops"
    ]

    # Store generated images data
    images_data = []

    # Generate images for each prompt
    for i, prompt in enumerate(prompts):
        print(f"Generating image {i+1}/{len(prompts)}: {prompt}")
        
        success, result = generate_image_with_nanobanana_pro(prompt, nano_banana_key)
        
        if success:
            print(f"  ✓ Successfully generated image: {result}")
            images_data.append((prompt, result))
        else:
            print(f"  ✗ Failed to generate image: {result}")
        
        print()

    # Save images to DOCX if any were generated
    if images_data:
        print("Saving generated images to DOCX file...")
        result = save_images_to_docx(images_data, "generated_images_report.docx")
        print(result)
        
        # Clean up temporary image files after saving to DOCX
        for _, image_path in images_data:
            try:
                os.remove(image_path)
                print(f"Cleaned up temporary file: {image_path}")
            except OSError:
                print(f"Could not remove temporary file: {image_path}")
    else:
        print("No images were generated successfully, so no DOCX file was created.")

    print("\n" + "=" * 70)
    print("Image generation and DOCX creation test completed.")


if __name__ == "__main__":
    main()