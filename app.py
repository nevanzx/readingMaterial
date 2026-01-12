import streamlit as st
import json
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.shared import OxmlElement, qn
import requests
import tempfile
from io import BytesIO
import re

# Function to add formatted text to a paragraph
def add_formatted_text(paragraph, text):
    """
    Adds text to a paragraph with formatting based on HTML-like tags.
    Supports <b> for bold and <i> for italic.
    """
    # Replace HTML entities to handle special characters
    text = text.replace('&lt;', '<').replace('&gt;', '>')

    # Split text by tags while keeping the tags
    parts = re.split(r'(<b>|</b>|<i>|</i>)', text)

    # Track current formatting state
    bold = False
    italic = False

    for part in parts:
        if part == '<b>':
            bold = True
        elif part == '</b>':
            bold = False
        elif part == '<i>':
            italic = True
        elif part == '</i>':
            italic = False
        elif part:  # Non-empty text content
            run = paragraph.add_run(part)
            run.font.name = 'Bookman Old Style'
            run.font.size = Pt(12)  # 12 points
            run.bold = bold
            run.italic = italic

# Function to generate image from prompt using an image generation API
def generate_image_from_prompt(prompt, ratio="1:1", size="4", api_key=None):
    """
    Generates an image from a prompt using the Gemini 3 Pro Image Preview model via the SDK.
    This function calls the Google's image generation API to generate the image.
    Returns the path to the downloaded image file, or None if the API call fails.
    """
    try:
        # Import the Google GenAI SDK
        try:
            from google import genai
            from google.genai import types
            GENAI_AVAILABLE = True
        except ImportError:
            print("google-genai not available. Please install it with: pip install google-genai")
            return None

        # Check if we have a valid API key to call the image generation service
        if not api_key or not api_key.strip() or "YOUR_" in api_key:
            print("No valid API key provided for image generation")
            return None

        # Attempt to call the Gemini 3 Pro Image Preview API with the prompt
        try:
            client = genai.Client(api_key=api_key)

            # Map ratio and size to appropriate inch dimensions
            ratio_map = {
                "16:9": (size, size * 9/16),  # For 16:9, if width is 'size', height is size * 9/16
                "4:3": (size, size * 3/4),    # For 4:3, if width is 'size', height is size * 3/4
                "3:4": (size * 3/4, size),    # For 3:4, if height is 'size', width is size * 3/4
                "1:1": (size, size)           # For 1:1, both dimensions are 'size'
            }

            # Get the mapped dimensions
            width_in, height_in = ratio_map.get(ratio, (4, 4))  # Default to 4x4 if ratio not recognized

            # Enhance prompt with ratio and resolution requirements
            full_prompt = f"{prompt}. Generate image with {ratio} aspect ratio at 1K resolution. Focus on maintaining the specified aspect ratio."

            # Call the API without the problematic response_mime_type config
            response = client.models.generate_content(
                model="gemini-3-pro-image-preview",
                contents=[full_prompt],
                # If you need safety settings, add them here, but do NOT add response_mime_type="image/jpeg"
            )

            # Check for image data in the response
            if response.candidates:
                for candidate in response.candidates:
                    if hasattr(candidate, 'content') and candidate.content.parts:
                        for part in candidate.content.parts:
                            # Check for inline_data (image)
                            if hasattr(part, 'inline_data') and part.inline_data:
                                # Create a temporary file
                                with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as temp_file:
                                    temp_file.write(part.inline_data.data)
                                    print(f"Successfully received and saved image from Gemini 3 Pro Image Preview for prompt: '{prompt[:50]}{'...' if len(prompt) > 50 else ''}'")
                                    return temp_file.name

            # If no image found, check for text feedback (e.g., safety refusal)
            if response.candidates and response.candidates[0].content.parts:
                text_part = response.candidates[0].content.parts[0].text
                if text_part:
                    print(f"Model returned text instead of image: {text_part[:100]}...")
                    return None

            print(f"No image data found in response for prompt: '{prompt}'")
            return None

        except Exception as e:
            print(f"SDK Error: {str(e)}")
            print(f"Failed to generate image for prompt: '{prompt}'")
            return None

    except Exception as e:
        print(f"Error preparing Gemini 3 Pro Image Preview API call: {e}")
        return None

# Function to add images to a paragraph with square text wrapping
def add_image_to_doc(doc, image_file, width=None, height=None, position='center'):
    """
    Adds an image to the document with specified positioning and square text wrapping.
    """
    # Create a temporary file to save the image stream
    if hasattr(image_file, 'name') and hasattr(image_file, 'getvalue'):
        # It's a file-like object from upload
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(image_file.name)[1]) as tmp_img:
            tmp_img.write(image_file.getvalue())
            tmp_img_path = tmp_img.name
    else:
        # It's a path to a generated image
        tmp_img_path = image_file

    try:
        # Add the image to the document
        paragraph = doc.add_paragraph()

        # Set alignment based on position
        if position == 'left':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif position == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Default center

        run = paragraph.add_run()
        picture = run.add_picture(tmp_img_path, width=width, height=height)

        # Get the picture element to modify its properties for square text wrapping
        pic = picture._inline
        # Set the distance from text - this affects the square wrapping appearance
        pic.dist_t = 0  # Distance from top
        pic.dist_b = 0  # Distance from bottom
        pic.dist_l = 114300  # Distance from left (in EMUs - 114300 EMUs = 0.1 inch)
        pic.dist_r = 114300  # Distance from right (in EMUs - 114300 EMUs = 0.1 inch)

        # Add some spacing after the image
        paragraph.paragraph_format.space_after = Inches(0.0833)  # 6/72 inches
    finally:
        # Clean up the temporary image file if it was created
        if hasattr(image_file, 'name') and hasattr(image_file, 'getvalue'):
            os.unlink(tmp_img_path)

# Set page config
st.set_page_config(
    page_title="College Reading Material Generator",
    page_icon="📚",
    layout="wide"
)

# App title
st.title("📚 College Reading Material Generator")
st.markdown("Generate reading materials for college students using AI")

# Initialize session state
if 'course_outline' not in st.session_state:
    st.session_state.course_outline = None
if 'chapters_data' not in st.session_state:
    st.session_state.chapters_data = None
if 'selected_chapters' not in st.session_state:
    st.session_state.selected_chapters = []
if 'generated_materials' not in st.session_state:
    st.session_state.generated_materials = {}
if 'processing_status' not in st.session_state:
    st.session_state.processing_status = ""
if 'generated_docs' not in st.session_state:
    st.session_state.generated_docs = {}  # Cache generated documents
if 'image_cache' not in st.session_state:
    st.session_state.image_cache = {}  # Cache generated images
if 'generation_complete' not in st.session_state:
    st.session_state.generation_complete = False  # Flag to indicate generation is done
if 'json_file_uploaded' not in st.session_state:
    st.session_state.json_file_uploaded = False  # Flag to indicate if JSON file was uploaded
if 'full_materials_json' not in st.session_state:
    st.session_state.full_materials_json = None  # Stores complete materials JSON if uploaded
if 'image_prompts_txt' not in st.session_state:
    st.session_state.image_prompts_txt = None  # Stores image prompts text for download
if 'original_filename' not in st.session_state:
    st.session_state.original_filename = None  # Stores original uploaded filename

# Sidebar for API keys
with st.sidebar:
    st.header("API Configuration")

    # Option to upload API keys from JSON file
    uploaded_keys = st.file_uploader("Upload API Keys (JSON)", type=["json"])

    if uploaded_keys:
        try:
            keys_data = json.load(uploaded_keys)
            apis = {api['name']: api['keys'][0] for api in keys_data['apis']}
            gemini_api_key = apis.get('gemini', '')
            nano_banana_gemini_api_key = apis.get('nano_banana_gemini', '')
            deepseek_api_key = apis.get('deepseek', '')
            st.success("API keys loaded from file!")
        except Exception as e:
            st.error(f"Error loading API keys: {str(e)}")
            gemini_api_key = st.text_input("Gemini API Key", type="password")
            nano_banana_gemini_api_key = st.text_input("Nano Banana Pro Gemini API Key", type="password")
            deepseek_api_key = st.text_input("Deepseek API Key", type="password")
    else:
        gemini_api_key = st.text_input("Gemini API Key", type="password")
        nano_banana_gemini_api_key = st.text_input("Nano Banana Pro Gemini API Key", type="password")
        deepseek_api_key = st.text_input("Deepseek API Key", type="password")

    # Toggle for saving image prompts to text file
    save_prompts_to_txt = st.toggle("Generate Image Prompts to Text File", value=False)

    # Show number of images to generate when toggle is enabled
    if save_prompts_to_txt:
        num_images_for_txt = st.number_input(
            "Number of images per chapter (for text file)",
            min_value=1,
            max_value=10,
            value=3,
            step=1
        )
    else:
        num_images_for_txt = 0

    # Toggle for image generation from Nano Banana Pro
    enable_image_generation = st.toggle("Enable Image Generation from Nano Banana Pro", value=False)

    # Show number of images to generate when toggle is enabled
    if enable_image_generation:
        num_images_per_chapter = st.number_input(
            "Number of images per chapter",
            min_value=1,
            max_value=10,
            value=3,
            step=1
        )
    else:
        num_images_per_chapter = 0

    # Context locality dropdown
    context_locality = st.selectbox(
        "Select context locality",
        options=["Philippines", "International", "Both"],
        index=0,  # Default to Philippines
        help="Choose the context for the generated content"
    )

    # Temperature setting for Deepseek
    deepseek_temperature = st.slider("Deepseek Temperature", min_value=0.0, max_value=1.0, value=0.6, step=0.1)

st.markdown("### Step 1: Upload Course Outline")
st.markdown("Upload a DOCX file containing your course outline and checklist or upload a previously saved JSON file")

# Tabs for choosing between DOCX upload or JSON upload
tab1, tab2 = st.tabs(["Upload DOCX File", "Upload Saved JSON File"])

with tab1:
    # File upload
    uploaded_file = st.file_uploader("Choose a DOCX file", type=["docx"])

    if uploaded_file is not None:
        st.session_state.course_outline = uploaded_file
        st.session_state.original_filename = os.path.splitext(uploaded_file.name)[0]  # Store base filename without extension
        st.success(f"File '{uploaded_file.name}' uploaded successfully!")

        # Display file info
        st.write(f"File size: {uploaded_file.size} bytes")

        # Note: Image generation is now handled automatically when Deepseek generates content with image prompts
        if enable_image_generation:
            st.info("Image generation is handled automatically when Deepseek generates content with image prompts")

        # Process the document
        if st.button("Process Course Outline with Gemini"):
            # Check for appropriate API key based on toggle
            if enable_image_generation and not nano_banana_gemini_api_key:
                st.error("Please enter your Nano Banana Pro Gemini API key in the sidebar")
            elif not enable_image_generation and not gemini_api_key:
                st.error("Please enter your Gemini API key in the sidebar")
            else:
                with st.spinner("Processing with Gemini 2.5..."):
                    # Extract text from the uploaded DOCX file
                    from docx import Document as DocxDocument
                    doc = DocxDocument(uploaded_file)

                    # Extract all text from the document
                    full_text = []
                    for paragraph in doc.paragraphs:
                        full_text.append(paragraph.text)

                    # Extract text from tables if any
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                full_text.append(cell.text)

                    course_outline_text = "\n".join(full_text)

                    # Call Gemini API to extract chapters and topics
                    try:
                        headers = {
                            "Content-Type": "application/json"
                        }

                        # Prepare the prompt for Gemini
                        prompt = f"""
                        Analyze the following course outline and extract chapters and topics in JSON format.
                        The JSON should have this required structure like this:
                        {{
                            "chapters": [
                                {{
                                    "chapter": "Chapter 1: Name",
                                    "topics": ["Topic 1", "Topic 2", "Topic 3"]
                                }}
                            ]
                        }}

                        Ensure that topics are specific and detailed enough to generate comprehensive reading materials.
                        Course Outline:
                        {course_outline_text}
                        """

                        # Use only the regular Gemini API key for outline generation (nano_banana is only for images)
                        api_key_to_use = gemini_api_key

                        # Make request to Gemini API
                        # For Gemini API, the key should be passed as a query parameter or in the URL
                        import urllib.parse
                        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={urllib.parse.quote(api_key_to_use)}"

                        response = requests.post(
                            url,
                            headers={"Content-Type": "application/json"},
                            json={
                                "contents": [{
                                    "parts": [{
                                        "text": prompt
                                    }]
                                }]
                            }
                        )

                        if response.status_code == 200:
                            result = response.json()

                            # Extract the text response
                            text_response = result['candidates'][0]['content']['parts'][0]['text']

                            # Find the JSON part in the response
                            import re
                            json_match = re.search(r'\{.*\}', text_response, re.DOTALL)

                            if json_match:
                                json_str = json_match.group()
                                st.session_state.chapters_data = json.loads(json_str)

                                # Clear selected chapters
                                st.session_state.selected_chapters = []

                                # Reset the flag since we're using Gemini
                                st.session_state.json_file_uploaded = False

                                st.success("Course outline processed successfully!")
                                st.session_state.processing_status = "processed"
                            else:
                                st.error("Could not extract JSON from Gemini response")
                        else:
                            st.error(f"Gemini API error: {response.status_code} - {response.text}")

                    except requests.exceptions.RequestException as e:
                        st.error(f"Network error while processing with Gemini: {str(e)}")
                    except json.JSONDecodeError:
                        st.error("Invalid JSON response from Gemini API")
                    except Exception as e:
                        st.error(f"Unexpected error while processing with Gemini: {str(e)}")

with tab2:
    # JSON file upload
    uploaded_json = st.file_uploader("Choose a saved JSON file", type=["json"])

    if uploaded_json is not None:
        try:
            # Load the JSON data
            json_data = json.load(uploaded_json)

            # Extract base filename from the uploaded JSON file name (without extension)
            json_base_filename = os.path.splitext(uploaded_json.name)[0]

            # Check if this is a generated materials JSON (has title, introduction, topics, summary, references)
            if all(key in json_data for key in ["title", "introduction", "topics", "summary", "references"]):
                # This is a generated materials JSON file with complete content
                # We need to convert it to the chapters_data format for the UI
                chapters_data = {
                    "chapters": []
                }

                # Extract chapter info from the material
                title = json_data.get("title", "Untitled Chapter")
                topics = []

                # Extract topic names from the material's topics
                for topic in json_data.get("topics", []):
                    topic_name = topic.get("topic", "Untitled Topic")
                    topics.append(topic_name)

                chapters_data["chapters"].append({
                    "chapter": title,
                    "topics": topics
                })

                st.session_state.chapters_data = chapters_data
                st.session_state.selected_chapters = []

                # Set the flag to indicate JSON file was uploaded
                st.session_state.json_file_uploaded = True
                # Store the complete material data in a separate variable to indicate it has full content
                st.session_state.full_materials_json = {title: json_data}
                # Store the JSON filename as the original filename for image prompts
                st.session_state.original_filename = json_base_filename

                st.success("Generated materials JSON file loaded successfully!")
                st.session_state.processing_status = "processed"
            elif "chapters" in json_data and isinstance(json_data["chapters"], list):
                # This is a chapter outline JSON file
                st.session_state.chapters_data = json_data
                st.session_state.selected_chapters = []

                # Set the flag to indicate JSON file was uploaded
                st.session_state.json_file_uploaded = True
                # Store the JSON filename as the original filename for image prompts
                st.session_state.original_filename = json_base_filename
                # Clear any previous full materials JSON
                if hasattr(st.session_state, 'full_materials_json'):
                    del st.session_state.full_materials_json

                st.success("Chapter outline JSON file loaded successfully!")
                st.session_state.processing_status = "processed"
            else:
                st.error("Invalid JSON format. The file should contain either a 'chapters' array or complete material data with 'title', 'introduction', etc.")
        except json.JSONDecodeError:
            st.error("Invalid JSON file format.")
        except Exception as e:
            st.error(f"Error loading JSON file: {str(e)}")

# Display chapters if available
if st.session_state.chapters_data:
    st.markdown("### Step 2: Select Chapters to Generate")

    # Add a button to save the JSON response
    if st.button("Save Chapters Data as JSON"):
        # Convert the chapters_data to JSON string (the outline structure)
        json_str = json.dumps(st.session_state.chapters_data, indent=2)

        # Create a download button for the JSON
        st.download_button(
            label="Download JSON File",
            data=json_str,
            file_name="chapters_data.json",
            mime="application/json"
        )
        st.info("Click the 'Download JSON File' button above to save the chapters data.")

    # Display chapters as checkboxes
    selected_chapters = []
    for idx, chapter in enumerate(st.session_state.chapters_data.get("chapters", [])):
        chapter_name = chapter.get("chapter", f"Chapter {idx+1}")
        topics = chapter.get("topics", [])

        col1, col2 = st.columns([1, 4])
        with col1:
            is_selected = st.checkbox(f"Include {chapter_name}", key=f"chapter_{idx}")
        with col2:
            st.write(f"**{chapter_name}**")
            if topics:
                st.write("Topics:")
                for topic in topics:
                    st.write(f"- {topic}")


        if is_selected:
            selected_chapters.append({
                "chapter": chapter_name,
                "topics": topics
            })

    # Update session state with selected chapters
    st.session_state.selected_chapters = selected_chapters

    # Generate materials for selected chapters
    if st.button("Generate Reading Materials") and selected_chapters:
        # Check if we're using a JSON file that contains complete materials (not just outline)
        using_complete_materials = (
            st.session_state.json_file_uploaded and
            hasattr(st.session_state, 'full_materials_json') and
            st.session_state.full_materials_json is not None
        )

        if not deepseek_api_key and not using_complete_materials:
            st.error("Please enter your Deepseek API key in the sidebar")
        else:
            # Clean up previously cached documents and reset flags
            for doc_path in st.session_state.generated_docs.values():
                try:
                    os.unlink(doc_path)
                except:
                    pass  # Ignore errors if file doesn't exist

            # Clean up cached images
            for img_path in st.session_state.image_cache.values():
                try:
                    os.unlink(img_path)
                except:
                    pass  # Ignore errors if file doesn't exist

            # Reset session state for new generation
            st.session_state.generated_docs = {}
            st.session_state.image_cache = {}
            st.session_state.generation_complete = False

            # Create a progress bar
            progress_bar = st.progress(0)
            status_text = st.empty()

            with st.spinner("Generating materials with Deepseek..."):
                # Process each selected chapter with Deepseek
                generated_materials = {}

                for idx, chapter_data in enumerate(selected_chapters):
                    chapter_name = chapter_data["chapter"]
                    topics = chapter_data["topics"]

                    status_text.text(f"Processing: {chapter_name}...")
                    progress_bar.progress((idx + 1) / len(selected_chapters))

                    # Prepare prompt for Deepseek Reasoner
                    base_prompt = f"""
                    You are an expert educator. Generate detailed, well-structured reading material for the following chapter and topics.
                    Use your reasoning capabilities to provide comprehensive explanations, examples, and connections between concepts.
                    ULTRA_CRITIAL: Return the response in JSON format with the following structure:
                    {{
                        "title": "Chapter Title",
                        "introduction": "Brief introduction to the chapter",
                        "topics": [
                            {{
                                "topic": "Topic Name",
                                "content": [
                                    "Content paragraph 1",
                                    "Content paragraph 2",
                                    "Content paragraph 3"
                                ]
                            }}
                        ],
                        "summary": "Comprehensive summary of the chapter content",
                        "references": [
                            {{
                                "title": "Title of source",
                                "author": "Author name(s)",
                                "year": "Publication year",
                                "publisher": "Publisher name",
                                "url": "URL if available"
                            }}
                        ]
                    }}

                    When generating content, use HTML-like tags to emphasize important words or concepts:
                    - Use <b> and </b> for bold formatting (for key terms, definitions, important concepts)
                    - Use <i> and </i> for italic formatting (for examples, clarifications, or special notes)

                    Context locality: {context_locality}
                    Chapter: {chapter_name}
                    Topics: {', '.join(topics)}

                    - ULTRA-CRITICAL: For each topic, generate 3-5 paragraphs depending on the complexity and length of explanation needed.
                    - Ensure the content is comprehensive, well-structured, and suitable for college-level students.
                    - Include a comprehensive summary section that captures the key points from all topics covered in the chapter.
                    - Include a references section with at least 3-5 scholarly sources in APA 7th edition format.
                    - For references in APA 7th edition format, follow these guidelines:
                      * Author, A. A. (Year). Title of work. Publisher. URL (if applicable)
                      * For journal articles: Author, A. A. (Year). Title of article. Title of Periodical, volume(issue), pages. https://doi.org/xx.xxx/yyyy
                      * For online sources: Author, A. A. (Year, Month Date). Title of webpage. Site Name. URL
                    - When context locality is 'Philippines' or 'Both', incorporate relevant examples, case studies, and references from the Philippines. ONLY MENTION THE WORD "PHILIPPINES" 2 times. The Philippine context must be implied.
                    - When context locality is 'International' or 'Both', include global examples and perspectives.
                    """

                    # If image generation is enabled, modify the prompt to include image generation
                    if enable_image_generation or save_prompts_to_txt:
                        # Select the specified number of topics for image generation based on which toggle is enabled
                        import random
                        if enable_image_generation:
                            num_images_to_generate = min(num_images_per_chapter, len(topics))
                        else:
                            num_images_to_generate = min(num_images_for_txt, len(topics))

                        selected_topics = random.sample(topics, num_images_to_generate)
                        selected_topics_str = ', '.join(selected_topics)

                        prompt = f"""
                        {base_prompt}

                        ADDITIONAL REQUIREMENTS FOR IMAGE GENERATION:
                        - For {num_images_to_generate} of the topics ({selected_topics_str}), also generate an image prompt that describes a relevant visual representation for the topic.
                        - Add an "image_prompt" field to the topic object for topics that should have images.
                        - Add a "ratio" field with one of these values: "16:9", "4:3", "3:4", or "1:1" based on what fits best for the image content.
                        - Add a "size" field with one of these values: "6" for 16:9, "3" for 4:3, "4" for 3:4, or "4" for 1:1 based on the chosen ratio.
                        - The image_prompt should be a detailed description that could be used to generate an appropriate educational image.
                        - Example topic structure with image prompt:
                        {{
                            "topic": "Topic Name",
                            "content": [
                                "Content paragraph 1",
                                "Content paragraph 2",
                                "Content paragraph 3"
                            ],
                            "image_prompt": "A detailed description of an image that visually represents this topic...",
                            "ratio": "16:9",
                            "size": "6"
                        }}
                        - Only add the image_prompt, ratio, and size fields to {num_images_to_generate} topics that would benefit most from a visual representation.
                        - Use 1K resolution for all generated images.
                        """
                    else:
                        prompt = base_prompt

                    try:
                        # Check if we're using a JSON file that contains complete materials (not just outline)
                        using_complete_materials = (
                            st.session_state.json_file_uploaded and
                            hasattr(st.session_state, 'full_materials_json') and
                            st.session_state.full_materials_json is not None and
                            chapter_name in st.session_state.full_materials_json
                        )

                        # If we have complete material data for this chapter, use it directly
                        if using_complete_materials:
                            material_data = st.session_state.full_materials_json[chapter_name]
                            generated_materials[chapter_name] = material_data
                            st.success(f"✓ Loaded material for: {chapter_name} from JSON file")
                        else:
                            # Call Deepseek API to generate content
                            headers = {
                                "Content-Type": "application/json",
                                "Authorization": f"Bearer {deepseek_api_key}"
                            }

                            response = requests.post(
                                "https://api.deepseek.com/chat/completions",
                                headers=headers,
                                json={
                                    "model": "deepseek-reasoner",
                                    "messages": [
                                        {"role": "user", "content": prompt}
                                    ],
                                    "temperature": deepseek_temperature,
                                    "max_tokens": 16000,  # Set token limit to 16k
                                    "response_format": {"type": "json_object"}
                                }
                            )

                            if response.status_code == 200:
                                result = response.json()
                                content = result['choices'][0]['message']['content']

                                # Extract JSON from response
                                import re
                                json_match = re.search(r'\{.*\}', content, re.DOTALL)

                                if json_match:
                                    json_str = json_match.group()
                                    material_data = json.loads(json_str)
                                    generated_materials[chapter_name] = material_data
                                    st.success(f"✓ Generated material for: {chapter_name}")
                                else:
                                    st.warning(f"Could not extract JSON for {chapter_name}")
                            else:
                                st.error(f"Deepseek API error for {chapter_name}: {response.status_code} - {response.text}")

                    except requests.exceptions.RequestException as e:
                        st.error(f"Network error while generating material for {chapter_name}: {str(e)}")
                    except json.JSONDecodeError:
                        st.error(f"Invalid JSON response from Deepseek API for {chapter_name}")
                    except Exception as e:
                        st.error(f"Unexpected error while generating material for {chapter_name}: {str(e)}")

                # If save_prompts_to_txt is enabled, save image prompts to a text file with the same name as the generated DOCX files
                if save_prompts_to_txt and generated_materials:
                    # Create a text file using the names of the selected chapters
                    # Use the first few chapter names to create a meaningful filename
                    chapter_names = list(generated_materials.keys())
                    if chapter_names:
                        # Take first chapter name or combine first few if many chapters are selected
                        if len(chapter_names) == 1:
                            base_filename = chapter_names[0].replace(' ', '_').replace(':', '_')
                        else:
                            # Use first chapter name followed by count of remaining chapters
                            first_chapter = chapter_names[0].replace(' ', '_').replace(':', '_')
                            base_filename = f"{first_chapter}_and_{len(chapter_names)-1}_more"
                    else:
                        # Fallback to original filename if no materials were generated
                        base_filename = st.session_state.original_filename or "generated_materials"

                    txt_filename = f"{base_filename}_image_prompts.txt"

                    # Collect all image prompts
                    all_prompts = []
                    for chapter_name, content in generated_materials.items():
                        for topic in content.get("topics", []):
                            image_prompt = topic.get("image_prompt")
                            if image_prompt:
                                all_prompts.append(f"Chapter: {chapter_name}\nTopic: {topic.get('topic', 'N/A')}\nPrompt: {image_prompt}\n---\n")

                    # Create the text content
                    txt_content = "Generated Image Prompts:\n\n" + "\n".join(all_prompts)

                    # Store the content in session state for download
                    st.session_state.image_prompts_txt = {
                        'content': txt_content,
                        'filename': txt_filename
                    }

                    st.success(f"Image prompts are ready for download")
                elif save_prompts_to_txt and not generated_materials:
                    st.warning("No materials were generated, so no image prompts to save.")

                st.session_state.generated_materials = generated_materials

                if generated_materials:
                    st.success("All reading materials generated successfully!")

                    # Add a button to save the generated materials as JSON
                    if st.button("Save Generated Materials as JSON"):
                        # Convert the generated materials to JSON string
                        json_str = json.dumps(generated_materials, indent=2)

                        # Create a download button for the JSON
                        st.download_button(
                            label="Download Generated Materials JSON",
                            data=json_str,
                            file_name="generated_materials.json",
                            mime="application/json"
                        )
                        st.info("Click the 'Download Generated Materials JSON' button above to save the complete materials data.")
                else:
                    st.warning("No materials were generated. Please check the API responses.")

            # Clear progress indicators
            progress_bar.empty()
            status_text.empty()

# Download section
if st.session_state.generated_materials:
    st.markdown("### Step 3: Download Generated Materials")

    # Check if documents are already cached
    if not st.session_state.generated_docs or not st.session_state.generation_complete:
        # Generate documents and cache them
        for chapter_name, content in st.session_state.generated_materials.items():
            # Check if document is already cached
            if chapter_name not in st.session_state.generated_docs:
                # Create DOCX file
                doc = Document()
                title_heading = doc.add_heading(content.get("title", chapter_name), 0)
                # Format the main title heading with 10/72 before line spacing and 6/72 after line spacing
                title_heading.paragraph_format.space_before = Inches(0.1389)  # 10/72 inches
                title_heading.paragraph_format.space_after = Inches(0.0833)   # 6/72 inches
                # Apply Bookman Old Style font to title
                for run in title_heading.runs:
                    run.font.name = 'Bookman Old Style'
                    run.font.size = Pt(12)  # 12 points

                # Add chapter introduction if available
                introduction = content.get("introduction")
                if introduction:
                    intro_heading = doc.add_heading("Chapter Introduction", level=1)
                    # Format the heading with 10/72 before line spacing and 6/72 after line spacing
                    intro_heading.paragraph_format.space_before = Inches(0.1389)  # 10/72 inches
                    intro_heading.paragraph_format.space_after = Inches(0.0833)   # 6/72 inches
                    # Apply Bookman Old Style font to heading
                    for run in intro_heading.runs:
                        run.font.name = 'Bookman Old Style'
                        run.font.size = Pt(12)  # 12 points

                    # Process introduction with formatting
                    intro_para = doc.add_paragraph()
                    # Format the paragraph: justified, 6/72 line spacing, 1 tab indent
                    intro_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    intro_para_format = intro_para.paragraph_format
                    intro_para_format.space_after = Inches(0.0833)  # 6/72 inches after paragraph
                    intro_para_format.line_spacing = 1.0  # Single line spacing
                    intro_para_format.first_line_indent = Inches(0.5)  # 0.5 inch indent
                    # Also set left indent to make it more visible
                    intro_para_format.left_indent = 0  # Keep left margin at 0

                    # Add formatted text to the paragraph
                    add_formatted_text(intro_para, introduction)

                for topic in content.get("topics", []):
                    # Add topic heading with 10pt before spacing and 6pt after spacing
                    topic_heading = doc.add_heading(topic.get("topic", "Topic"), level=1)
                    topic_heading.paragraph_format.space_before = Inches(0.1389)  # Approximately 10pt (10/72 inches)
                    topic_heading.paragraph_format.space_after = Inches(0.0833)   # Approximately 6pt (6/72 inches)
                    # Apply Bookman Old Style font to heading
                    for run in topic_heading.runs:
                        run.font.name = 'Bookman Old Style'
                        run.font.size = Pt(12)  # 12 points

                    # Handle both single content string and multiple content sections
                    topic_content = topic.get("content", "")

                    # Check if this topic has an image prompt and if image generation is enabled
                    image_prompt = topic.get("image_prompt")
                    image_ratio = topic.get("ratio", "1:1")  # Default to square if no ratio specified
                    image_size = float(topic.get("size", "4"))  # Default to 4 inches if no size specified
                    image_generated = None

                    # Only generate images if the nano banana image generation toggle is enabled
                    if enable_image_generation and image_prompt:
                        # Create a unique key for this image based on the prompt
                        image_key = hash(image_prompt)

                        # Check if image is already cached
                        if image_key in st.session_state.image_cache:
                            image_generated = st.session_state.image_cache[image_key]
                        else:
                            # Generate image from prompt using Nano Banana Pro Gemini API with ratio and size
                            image_generated = generate_image_from_prompt(image_prompt, image_ratio, image_size, nano_banana_gemini_api_key)
                            # Cache the generated image path
                            if image_generated:
                                st.session_state.image_cache[image_key] = image_generated

                    if isinstance(topic_content, list):
                        # Multiple content sections
                        paragraph_count = 0
                        for content_section_idx, content_section in enumerate(topic_content):
                            para = doc.add_paragraph()
                            # Format the paragraph: justified, 6/72 after spacing, 1 tab indent
                            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            para_format = para.paragraph_format
                            para_format.space_after = Inches(0.0833)  # 6/72 inches after paragraph
                            para_format.line_spacing = 1.0  # Single line spacing
                            para_format.first_line_indent = Inches(0.5)  # 0.5 inch indent
                            # Also set left indent to make it more visible
                            para_format.left_indent = 0  # Keep left margin at 0

                            # Add formatted text to the paragraph
                            add_formatted_text(para, content_section)

                            # Add image after the second paragraph if image generation is enabled and image exists
                            paragraph_count += 1
                            if (enable_image_generation and image_generated and
                                paragraph_count == 2):  # After the second paragraph
                                # Determine position (left or right) randomly for variety
                                import random
                                position = random.choice(['left', 'right'])
                                # Use standardized dimensions based on the ratio and size from the topic
                                ratio = topic.get("ratio", "1:1")
                                size = float(topic.get("size", "4"))

                                # Calculate dimensions based on ratio
                                if ratio == "16:9":
                                    width_in, height_in = size, size * 9/16
                                elif ratio == "4:3":
                                    width_in, height_in = size, size * 3/4
                                elif ratio == "3:4":
                                    width_in, height_in = size * 3/4, size
                                elif ratio == "1:1":  # Square
                                    width_in, height_in = size, size
                                else:  # Default to square
                                    width_in, height_in = size, size

                                # Add image with standardized dimensions based on ratio, random placement
                                add_image_to_doc(doc, image_generated,
                                               width=Inches(width_in), height=Inches(height_in),
                                               position=position)
                    else:
                        # Single content string
                        para = doc.add_paragraph()
                        # Format the paragraph: justified, 6/72 after spacing, 1 tab indent
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        para_format = para.paragraph_format
                        para_format.space_after = Inches(0.0833)  # 6/72 inches after paragraph
                        para_format.line_spacing = 1.0  # Single line spacing
                        para_format.first_line_indent = Inches(0.5)  # 0.5 inch indent

                        # Add formatted text to the paragraph
                        add_formatted_text(para, topic_content)

                        # If there's only one paragraph and we have an image, add it after
                        if enable_image_generation and image_generated:
                            # Determine position (left or right) randomly for variety
                            import random
                            position = random.choice(['left', 'right'])
                            # Use standardized dimensions based on the ratio and size from the topic
                            ratio = topic.get("ratio", "1:1")
                            size = float(topic.get("size", "4"))

                            # Calculate dimensions based on ratio
                            if ratio == "16:9":
                                width_in, height_in = size, size * 9/16
                            elif ratio == "4:3":
                                width_in, height_in = size, size * 3/4
                            elif ratio == "3:4":
                                width_in, height_in = size * 3/4, size
                            elif ratio == "1:1":  # Square
                                width_in, height_in = size, size
                            else:  # Default to square
                                width_in, height_in = size, size

                            # Add image with standardized dimensions based on ratio, random placement
                            add_image_to_doc(doc, image_generated,
                                           width=Inches(width_in), height=Inches(height_in),
                                           position=position)

                # Add Summary section
                summary = content.get("summary")
                if summary:
                    summary_heading = doc.add_heading("Summary", level=1)
                    summary_heading.paragraph_format.space_before = Inches(0.1389)  # Approximately 10pt (10/72 inches)
                    summary_heading.paragraph_format.space_after = Inches(0.0833)   # Approximately 6pt (6/72 inches)
                    # Apply Bookman Old Style font to heading
                    for run in summary_heading.runs:
                        run.font.name = 'Bookman Old Style'
                        run.font.size = Pt(12)  # 12 points

                    summary_para = doc.add_paragraph()
                    # Format the paragraph: justified, 6/72 after spacing, 1 tab indent
                    summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    summary_para_format = summary_para.paragraph_format
                    summary_para_format.space_after = Inches(0.0833)  # 6/72 inches after paragraph
                    summary_para_format.line_spacing = 1.0  # Single line spacing
                    summary_para_format.first_line_indent = Inches(0.5)  # 0.5 inch indent

                    # Add formatted text to the paragraph
                    add_formatted_text(summary_para, summary)

                # Add References section
                references = content.get("references", [])
                if references:
                    references_heading = doc.add_heading("References", level=1)
                    references_heading.paragraph_format.space_before = Inches(0.1389)  # Approximately 10pt (10/72 inches)
                    references_heading.paragraph_format.space_after = Inches(0.0833)   # Approximately 6pt (6/72 inches)
                    # Apply Bookman Old Style font to heading
                    for run in references_heading.runs:
                        run.font.name = 'Bookman Old Style'
                        run.font.size = Pt(12)  # 12 points

                    for reference in references:
                        ref_para = doc.add_paragraph()
                        # Format the paragraph: justified, 6/72 after spacing, hanging indent of 0.5 inches
                        ref_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        ref_para_format = ref_para.paragraph_format
                        ref_para_format.space_after = Inches(0.0833)  # 6/72 inches after paragraph
                        ref_para_format.line_spacing = 1.0  # Single line spacing
                        # APA style uses a hanging indent of 0.5 inches
                        ref_para_format.first_line_indent = Inches(-0.5)
                        ref_para_format.left_indent = Inches(0.5)

                        # Format reference according to APA 7th edition
                        author = reference.get("author", "")
                        year = reference.get("year", "")
                        title = reference.get("title", "")
                        publisher = reference.get("publisher", "")
                        url = reference.get("url", "")

                        # Construct APA 7th edition reference
                        apa_reference = ""
                        if author:
                            apa_reference += f"{author}. "
                        if year:
                            apa_reference += f"({year}). "
                        if title:
                            apa_reference += f"<i>{title}</i>. "
                        if publisher:
                            apa_reference += f"{publisher}."
                        if url:
                            apa_reference += f" {url}"

                        # Add formatted text to the paragraph
                        add_formatted_text(ref_para, apa_reference)

                # Save to temporary file and cache it
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
                    doc.save(tmp_file.name)
                    st.session_state.generated_docs[chapter_name] = tmp_file.name  # Cache the file path

        # Mark generation as complete to prevent regeneration on rerun
        st.session_state.generation_complete = True

    # Display download buttons for cached documents
    for chapter_name in st.session_state.generated_materials.keys():
        st.subheader(f"{chapter_name}")

        if chapter_name in st.session_state.generated_docs:
            # Provide download button for cached document
            with open(st.session_state.generated_docs[chapter_name], "rb") as f:
                st.download_button(
                    label=f"Download {chapter_name}.docx",
                    data=f.read(),
                    file_name=f"{chapter_name.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_{chapter_name}"  # Unique key to prevent conflicts
                )

    # Show download button for image prompts if they were generated
    if st.session_state.image_prompts_txt:
        st.download_button(
            label="Download Image Prompts Text File",
            data=st.session_state.image_prompts_txt['content'],
            file_name=st.session_state.image_prompts_txt['filename'],
            mime="text/plain"
        )
        st.info(f"You can download the image prompts text file: {st.session_state.image_prompts_txt['filename']}")